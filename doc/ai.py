from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 문서 객체 생성
doc = Document()

# 스타일 설정 (기본 폰트 등은 시스템 종속적이므로 기본값 사용하되 크기 조정)
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic' # 워드에서 열 때 맑은 고딕 등으로 보이도록 설정 시도
style.font.size = Pt(11)

# 1. 표지
head = doc.add_heading('발표 시나리오 및 핸드아웃', 0)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[주제: 개발자의 진화 - Know-How에서 AI Design으로]')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph('\n')
doc.add_paragraph('작성일: 2025년 8월 26일')
doc.add_paragraph('발표자: (성함)')
doc.add_page_break()

# 2. 발표 개요
doc.add_heading('1. 발표 개요', level=1)
doc.add_paragraph('■ 목적: AI 시대에 변화하는 개발자의 핵심 역량(설계 능력)을 강조하고 동기 부여')
doc.add_paragraph('■ 핵심 메시지: "구현은 AI에게, 설계는 인간에게. 코더(Coder)에서 아키텍트(Architect)로 진화하라."')
doc.add_paragraph('■ 예상 소요 시간: 15~20분')
doc.add_paragraph('\n')

# 3. 상세 스크립트
doc.add_heading('2. 슬라이드별 상세 대본', level=1)

slides = [
    {
        "title": "Slide 1: 표지",
        "visual": "[텍스트] 개발자의 진화: Know-How → Know-Where → AI Design",
        "script": "안녕하세요. 오늘 저는 개발자라는 직업이 어떻게 변화해왔고, 앞으로 우리는 무엇을 준비해야 하는지에 대해 이야기하려 합니다. 특히 'AI 디자인'이라는 새로운 키워드를 중심으로 말이죠."
    },
    {
        "title": "Slide 2: 화두 던지기",
        "visual": "[텍스트] '코딩 잘한다'의 기준 변화 (암기 → 정보 → ???)",
        "script": "여러분은 '개발을 잘한다'는 것이 무엇이라고 생각하시나요? 과거에는 복잡한 알고리즘을 안 보고 짜는 것이 실력이었습니다. 조금 전까지만 해도 검색을 잘해서 남들이 짠 코드를 잘 활용하는 게 실력이었죠. 그렇다면, AI가 코드를 짜주는 지금, 과연 무엇이 실력일까요?"
    },
    {
        "title": "Slide 3: 변화의 흐름 (Roadmap)",
        "visual": "[도표] 1. Know-How(과거) → 2. Know-Where(현재) → 3. AI Design(미래)",
        "script": "저는 개발의 역사를 크게 세 가지 시대로 구분합니다. '어떻게'가 중요했던 노하우의 시대, '어디에'가 중요했던 노우웨어의 시대, 그리고 이제 도래한 'AI 디자인'의 시대입니다."
    },
    {
        "title": "Slide 4: Era 1. Know-How (노하우)",
        "visual": "[이미지] 쌓여있는 전공 서적과 터미널 / [키워드] 장인, 암기, 최적화",
        "script": "첫 번째는 '노하우'의 시대입니다. 인터넷이 느리거나 없던 시절, 개발자는 걸어 다니는 백과사전이어야 했습니다. 문법을 외우고, 최적화 기법을 손끝에 익히는 '장인 정신'이 핵심이었죠. 이때의 개발자는 코드를 한 줄 한 줄 깎아 만드는 조각가와 같았습니다."
    },
    {
        "title": "Slide 5: Era 2. Know-Where (노우웨어)",
        "visual": "[이미지] 구글 검색창과 스택오버플로우 / [키워드] 검색가, 조립, 오픈소스",
        "script": "두 번째는 우리가 익숙한 '노우웨어'의 시대입니다. 모든 것을 외울 필요가 없어졌습니다. 구글링을 잘하고, 적절한 오픈소스를 찾아내 조립하는 능력이 중요해졌죠. '어디에 답이 있는지 아는 것'이 곧 실력이었습니다. 개발자는 조각가에서 '정보를 조립하는 조립가'로 변모했습니다."
    },
    {
        "title": "Slide 6: Era 3. AI Design (AI 디자인)",
        "visual": "[이미지] AI 에이전트를 지휘하는 모습 / [키워드] 설계자, 지휘, 판단",
        "script": "그리고 이제, 세 번째 시대인 'AI 디자인'의 시대가 왔습니다. 코드를 짜는 행위(구현) 자체는 AI가 순식간에 처리합니다. 이제 개발자에게 필요한 건 '어떻게 짤까'가 아니라, '무엇을, 왜, 어떤 구조로 만들 것인가'를 결정하는 설계(Design) 능력입니다."
    },
    {
        "title": "Slide 7: 왜 'Design' 인가?",
        "visual": "[텍스트] Coding is Free, Thinking is Expensive. (Writer → Editor-in-Chief)",
        "script": "왜 제가 '디자인'이라고 표현했을까요? AI는 훌륭한 '작가'지만, 방향을 잡아줄 '편집장'이 필요하기 때문입니다. AI가 쏟아내는 코드가 우리 시스템에 맞는지, 보안엔 문제가 없는지 판단하고 전체 그림을 그리는 것. 그것이 바로 디자인입니다."
    },
    {
        "title": "Slide 8: 시대별 비교 요약",
        "visual": "[표] 핵심 역량 / 역할 / 작업 방식 / 생산성 비교",
        "script": "정리하자면 이렇습니다. 과거의 우리가 손으로 일하는 장인이었고, 현재가 정보를 찾는 사서였다면, 미래의 우리는 AI라는 유능한 건설팀을 지휘하는 '건축가(Architect)'가 되어야 합니다."
    },
    {
        "title": "Slide 9: 새로운 핵심 역량",
        "visual": "[텍스트] 1.질문 능력(Prompt) 2.안목(Insight) 3.시스템 사고(System Thinking)",
        "script": "그렇다면 우리는 무엇을 준비해야 할까요? 첫째, 정확히 지시하는 질문 능력. 둘째, AI가 짠 코드가 맞는지 가려낼 안목. 마지막으로, 숲을 보는 설계적 사고가 필수적입니다."
    },
    {
        "title": "Slide 10: 결론",
        "visual": "[텍스트] AI는 경쟁자가 아닌 비서입니다. Be the Director.",
        "script": "AI는 경쟁자가 아닙니다. 우리가 부려야 할 도구이자 비서입니다. 과거의 노하우와 노우웨어를 발판 삼아, 이제는 AI라는 강력한 엔진을 통해 여러분만의 소프트웨어를 '디자인' 하십시오."
    },
    {
        "title": "Slide 11: Q&A",
        "visual": "[텍스트] Thank You / 질의응답",
        "script": "경청해 주셔서 감사합니다. 질문 있으시면 편하게 말씀해 주세요."
    }
]

for slide in slides:
    doc.add_heading(slide["title"], level=2)
    p_visual = doc.add_paragraph()
    p_visual.add_run('화면 구성: ').bold = True
    p_visual.add_run(slide["visual"])
    
    p_script = doc.add_paragraph()
    p_script.add_run('🗣 대본: ').bold = True
    p_script.add_run(slide["script"])
    doc.add_paragraph('\n')

doc.add_page_break()

# 4. 배포용 요약표 (별첨)
doc.add_heading('3. [별첨] 시대별 개발 패러다임 변화 요약', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '구분'
hdr_cells[1].text = 'Know-How (과거)'
hdr_cells[2].text = 'Know-Where (현재)'
hdr_cells[3].text = 'AI Design (미래)'

data = [
    ['핵심 역량', '문법 암기, 최적화', '검색, 오픈소스 활용', '설계, 지휘, 검증'],
    ['개발자 역할', '장인 (Writer)', '조립가 (Assembler)', '감독 (Director)'],
    ['작업 방식', 'Typing (직접 입력)', 'Searching (검색)', 'Designing (설계)'],
    ['비유', '도공 (Craftsman)', '사서 (Librarian)', '건축가 (Architect)']
]

for item in data:
    row_cells = table.add_row().cells
    for i in range(4):
        row_cells[i].text = item[i]

doc.add_paragraph('\n* 이 표는 발표 내용의 핵심 요약으로, 청중 배포용 자료로 활용하실 수 있습니다.')

# 저장
file_path_docx = "Developer_Evolution_Script.docx"
doc.save(file_path_docx)
file_path_docx